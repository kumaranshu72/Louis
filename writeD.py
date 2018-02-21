from docx import Document
from docx.shared import Inches

class writeDoc:
    def __init__(self):
        self.document = Document()

    def addHeading(self,heading):
        self.document.add_heading(heading, 0)

    def addParagraph(self,paragraph):
        p = self.document.add_paragraph(paragraph)

    def addImage(self,img_path):
        self.document.add_picture(img_path, width=Inches(1.25))

    def addPage_break(self):
        self.document.add_page_break()
        pass

    def saveDoc(self,name):
        self.document.save(name)
        #self.document.cleanup()

'''
d1=writeDoc()
d1.addHeading("First Doc")
d1.addParagraph("My doc")
d1.saveDoc("Anshu.docx")

d1=writeDoc()
d1.addHeading("Second Doc")
d1.addParagraph("My doc")
d1.saveDoc("Anshu1.docx")
'''

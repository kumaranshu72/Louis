from docx import Document
from docx.shared import Inches

document = Document()

def add_heading(heading):
    document.add_heading(heading, 0)

def add_paragraph(paragraph):
    p = document.add_paragraph(paragraph)

def add_image(img_path):
    document.add_picture(img_path, width=Inches(1.25))

def add_page_break():
    document.add_page_break()
    pass

def save_doc(name):
    document.save(name)

add_heading("My First Doucument")
add_paragraph("FIrst One")
add_paragraph("second One")
save_doc("anshu.docx")


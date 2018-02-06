#import urllib.request
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

#urllib.request.urlretrieve("http://placehold.it/350x150", "placeholder.png")
#document.add_picture('placeholder.png', width=Inches(1.25))

recordset = [
    {
        "id" : 1,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 2,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 3,
        "qty": 2,
        "desc": "New item"
    },

]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item["qty"])
    row_cells[1].text = str(item["id"])
    row_cells[2].text = item["desc"]

document.add_page_break()

document.save('demo.docx')
